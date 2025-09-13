"""
Enhanced report export functionality for DRC climate action reporting platform.
Supports PDF and Word document generation with logo, signature, and chart inclusion.
"""

import os
import io
import json
from datetime import datetime
from typing import Dict, List, Optional, Any
import base64

# Import for document generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.linecharts import HorizontalLineChart
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

class ReportExporter:
    """Enhanced report exporter with logo, signature, and chart support."""
    
    def __init__(self, logo_path: str = 'logo.png', signature_path: str = 'signature.png'):
        """
        Initialize the report exporter.
        
        Args:
            logo_path: Path to the ministry logo
            signature_path: Path to the authorized signature
        """
        self.logo_path = logo_path
        self.signature_path = signature_path
        self.export_dir = 'exports'
        
        # Ensure export directory exists
        os.makedirs(self.export_dir, exist_ok=True)
        
        # Check for logo and signature files
        if not os.path.exists(self.logo_path):
            print(f"Warning: Logo file not found at {self.logo_path}")
        if not os.path.exists(self.signature_path):
            print(f"Warning: Signature file not found at {self.signature_path}")
    
    def create_chart_image(self, chart_data: Dict, chart_type: str = 'line') -> Optional[str]:
        """
        Create a chart image from data.
        
        Args:
            chart_data: Dictionary containing chart data
            chart_type: Type of chart ('line', 'bar', 'pie')
            
        Returns:
            Path to the generated chart image or None if failed
        """
        if not MATPLOTLIB_AVAILABLE:
            print("Warning: Matplotlib not available for chart generation")
            return None
        
        try:
            plt.figure(figsize=(10, 6))
            
            if chart_type == 'line':
                x_data = chart_data.get('x_data', [])
                y_data = chart_data.get('y_data', [])
                plt.plot(x_data, y_data, marker='o', linewidth=2, markersize=6)
                plt.xlabel(chart_data.get('x_label', 'X Axis'))
                plt.ylabel(chart_data.get('y_label', 'Y Axis'))
                
            elif chart_type == 'bar':
                x_data = chart_data.get('x_data', [])
                y_data = chart_data.get('y_data', [])
                plt.bar(x_data, y_data, color='steelblue', alpha=0.7)
                plt.xlabel(chart_data.get('x_label', 'Categories'))
                plt.ylabel(chart_data.get('y_label', 'Values'))
                plt.xticks(rotation=45, ha='right')
                
            elif chart_type == 'pie':
                labels = chart_data.get('labels', [])
                sizes = chart_data.get('sizes', [])
                plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
                plt.axis('equal')
            
            plt.title(chart_data.get('title', 'Chart'), fontsize=14, fontweight='bold')
            if chart_type != 'pie':
                plt.grid(True, alpha=0.3)
            
            # Save chart
            chart_filename = f"chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            chart_path = os.path.join(self.export_dir, chart_filename)
            plt.tight_layout()
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            # Verify file was created
            if os.path.exists(chart_path):
                return chart_path
            else:
                print(f"Chart file was not created at {chart_path}")
                return None
            
        except Exception as e:
            print(f"Error creating chart: {e}")
            plt.close()  # Make sure to close figure even on error
            return None
    
    def export_to_pdf(self, report_data: Dict, include_charts: bool = True, 
                     include_logo: bool = True, include_signature: bool = True) -> str:
        """
        Export report to PDF format.
        
        Args:
            report_data: Dictionary containing report data
            include_charts: Whether to include charts
            include_logo: Whether to include ministry logo
            include_signature: Whether to include signature
            
        Returns:
            Path to the generated PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export. Install with: pip install reportlab")
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"report_{report_data.get('title', 'untitled').replace(' ', '_')}_{timestamp}.pdf"
        filepath = os.path.join(self.export_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.darkblue
        )
        
        # Add logo
        if include_logo and os.path.exists(self.logo_path):
            try:
                logo = Image(self.logo_path, width=2*inch, height=1*inch)
                logo.hAlign = 'CENTER'
                story.append(logo)
                story.append(Spacer(1, 20))
            except Exception as e:
                print(f"Warning: Could not include logo: {e}")
        
        # Add title and header info
        story.append(Paragraph("Democratic Republic of Congo", styles['Normal']))
        story.append(Paragraph("Ministry of Environment and Sustainable Development", styles['Normal']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(report_data.get('title', 'Climate Action Report'), title_style))
        story.append(Spacer(1, 20))
        
        # Report metadata
        metadata_data = [
            ['Report Type:', report_data.get('report_type', 'N/A')],
            ['Province:', report_data.get('province', 'All Provinces')],
            ['Generated On:', datetime.now().strftime('%Y-%m-%d %H:%M')],
            ['Generated By:', report_data.get('author', 'System')],
            ['Status:', report_data.get('status', 'Draft')]
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 30))
        
        # Add description
        if report_data.get('description'):
            story.append(Paragraph("Executive Summary", header_style))
            story.append(Paragraph(report_data['description'], styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Add satellite data section
        if report_data.get('satellite_data'):
            story.append(Paragraph("Satellite Data Analysis", header_style))
            
            satellite_data = report_data['satellite_data']
            if satellite_data.get('forest_loss_data'):
                story.append(Paragraph("Forest Loss Analysis:", styles['Heading3']))
                
                # Create table for forest loss data
                forest_data = satellite_data['forest_loss_data']
                if isinstance(forest_data, list) and forest_data:
                    table_data = [['Year', 'Forest Loss (hectares)']]
                    for item in forest_data:
                        if isinstance(item, dict):
                            table_data.append([str(item.get('year', 'N/A')), 
                                             f"{item.get('loss_area_ha', 0):,.2f}"])
                    
                    forest_table = Table(table_data, colWidths=[1.5*inch, 2*inch])
                    forest_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    
                    story.append(forest_table)
                    story.append(Spacer(1, 20))
        
        # Add charts if requested
        if include_charts and report_data.get('charts'):
            story.append(Paragraph("Data Visualization", header_style))
            
            for chart_config in report_data['charts']:
                chart_path = self.create_chart_image(chart_config, chart_config.get('type', 'line'))
                if chart_path and os.path.exists(chart_path):
                    try:
                        chart_img = Image(chart_path, width=6*inch, height=3.6*inch)
                        chart_img.hAlign = 'CENTER'
                        story.append(chart_img)
                        story.append(Spacer(1, 20))
                        
                        # Clean up temporary chart file
                        os.remove(chart_path)
                    except Exception as e:
                        print(f"Warning: Could not include chart: {e}")
        
        # Add signature section
        if include_signature:
            story.append(Spacer(1, 40))
            story.append(Paragraph("Authorization", header_style))
            story.append(Spacer(1, 20))
            
            if os.path.exists(self.signature_path):
                try:
                    signature = Image(self.signature_path, width=2*inch, height=0.8*inch)
                    signature.hAlign = 'LEFT'
                    story.append(signature)
                except Exception as e:
                    print(f"Warning: Could not include signature: {e}")
            
            story.append(Spacer(1, 10))
            story.append(Paragraph("_" * 30, styles['Normal']))
            story.append(Paragraph("Authorized Ministry Official", styles['Normal']))
            story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        return filepath
    
    def export_to_word(self, report_data: Dict, include_charts: bool = True,
                      include_logo: bool = True, include_signature: bool = True) -> str:
        """
        Export report to Word document format.
        
        Args:
            report_data: Dictionary containing report data
            include_charts: Whether to include charts
            include_logo: Whether to include ministry logo
            include_signature: Whether to include signature
            
        Returns:
            Path to the generated Word document
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"report_{report_data.get('title', 'untitled').replace(' ', '_')}_{timestamp}.docx"
        filepath = os.path.join(self.export_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = report_data.get('title', 'Climate Action Report')
        doc.core_properties.author = report_data.get('author', 'DRC Ministry of Environment')
        doc.core_properties.subject = 'Climate Action Report'
        
        # Add logo
        if include_logo and os.path.exists(self.logo_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(2))
            except Exception as e:
                print(f"Warning: Could not include logo in Word document: {e}")
        
        # Add header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run("Democratic Republic of Congo\nMinistry of Environment and Sustainable Development")
        header_run.font.size = Pt(12)
        
        # Add title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(report_data.get('title', 'Climate Action Report'))
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Add metadata table
        metadata_table = doc.add_table(rows=5, cols=2)
        metadata_table.style = 'Table Grid'
        
        metadata_items = [
            ('Report Type:', report_data.get('report_type', 'N/A')),
            ('Province:', report_data.get('province', 'All Provinces')),
            ('Generated On:', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Generated By:', report_data.get('author', 'System')),
            ('Status:', report_data.get('status', 'Draft'))
        ]
        
        for i, (label, value) in enumerate(metadata_items):
            metadata_table.cell(i, 0).text = label
            metadata_table.cell(i, 1).text = value
        
        doc.add_paragraph()  # Empty line
        
        # Add description
        if report_data.get('description'):
            desc_heading = doc.add_paragraph()
            desc_heading_run = desc_heading.add_run("Executive Summary")
            desc_heading_run.font.size = Pt(14)
            desc_heading_run.font.bold = True
            
            doc.add_paragraph(report_data['description'])
            doc.add_paragraph()  # Empty line
        
        # Add satellite data section
        if report_data.get('satellite_data'):
            sat_heading = doc.add_paragraph()
            sat_heading_run = sat_heading.add_run("Satellite Data Analysis")
            sat_heading_run.font.size = Pt(14)
            sat_heading_run.font.bold = True
            
            satellite_data = report_data['satellite_data']
            if satellite_data.get('forest_loss_data'):
                forest_heading = doc.add_paragraph()
                forest_heading_run = forest_heading.add_run("Forest Loss Analysis:")
                forest_heading_run.font.bold = True
                
                # Create table for forest loss data
                forest_data = satellite_data['forest_loss_data']
                if isinstance(forest_data, list) and forest_data:
                    forest_table = doc.add_table(rows=len(forest_data) + 1, cols=2)
                    forest_table.style = 'Table Grid'
                    
                    # Header row
                    forest_table.cell(0, 0).text = 'Year'
                    forest_table.cell(0, 1).text = 'Forest Loss (hectares)'
                    
                    # Data rows
                    for i, item in enumerate(forest_data, 1):
                        if isinstance(item, dict):
                            forest_table.cell(i, 0).text = str(item.get('year', 'N/A'))
                            forest_table.cell(i, 1).text = f"{item.get('loss_area_ha', 0):,.2f}"
                
                doc.add_paragraph()  # Empty line
        
        # Add charts if requested
        if include_charts and report_data.get('charts'):
            chart_heading = doc.add_paragraph()
            chart_heading_run = chart_heading.add_run("Data Visualization")
            chart_heading_run.font.size = Pt(14)
            chart_heading_run.font.bold = True
            
            for chart_config in report_data['charts']:
                chart_path = self.create_chart_image(chart_config, chart_config.get('type', 'line'))
                if chart_path and os.path.exists(chart_path):
                    try:
                        chart_p = doc.add_paragraph()
                        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        chart_run = chart_p.add_run()
                        chart_run.add_picture(chart_path, width=Inches(6))
                        
                        # Clean up temporary chart file
                        os.remove(chart_path)
                    except Exception as e:
                        print(f"Warning: Could not include chart in Word document: {e}")
                
                doc.add_paragraph()  # Empty line
        
        # Add signature section
        if include_signature:
            doc.add_page_break()  # New page for signature
            
            sig_heading = doc.add_paragraph()
            sig_heading_run = sig_heading.add_run("Authorization")
            sig_heading_run.font.size = Pt(14)
            sig_heading_run.font.bold = True
            
            doc.add_paragraph()  # Empty line
            
            if os.path.exists(self.signature_path):
                try:
                    sig_p = doc.add_paragraph()
                    sig_run = sig_p.add_run()
                    sig_run.add_picture(self.signature_path, width=Inches(2))
                except Exception as e:
                    print(f"Warning: Could not include signature in Word document: {e}")
            
            doc.add_paragraph("_" * 30)
            doc.add_paragraph("Authorized Ministry Official")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Save document
        doc.save(filepath)
        return filepath
    
    def export_report(self, report_data: Dict, export_format: str = 'pdf', 
                     include_charts: bool = True, include_logo: bool = True, 
                     include_signature: bool = True) -> str:
        """
        Export report in the specified format.
        
        Args:
            report_data: Dictionary containing report data
            export_format: Format for export ('pdf' or 'docx')
            include_charts: Whether to include charts
            include_logo: Whether to include ministry logo
            include_signature: Whether to include signature
            
        Returns:
            Path to the generated file
        """
        if export_format.lower() == 'pdf':
            return self.export_to_pdf(report_data, include_charts, include_logo, include_signature)
        elif export_format.lower() in ['docx', 'word']:
            return self.export_to_word(report_data, include_charts, include_logo, include_signature)
        else:
            raise ValueError(f"Unsupported export format: {export_format}")

# Sample usage and testing
def create_sample_report_data():
    """Create sample report data for testing."""
    return {
        'title': 'Forest Loss Assessment for Kinshasa Province',
        'report_type': 'forest_loss',
        'province': 'Kinshasa',
        'status': 'Draft',
        'author': 'DRC Climate Team',
        'description': 'This report provides a comprehensive analysis of forest loss in Kinshasa province from 2020 to 2023, based on satellite data from Global Forest Watch. The analysis reveals significant trends in deforestation patterns and provides recommendations for conservation efforts.',
        'satellite_data': {
            'forest_loss_data': [
                {'year': 2020, 'loss_area_ha': 1250.5},
                {'year': 2021, 'loss_area_ha': 1387.2},
                {'year': 2022, 'loss_area_ha': 1456.8},
                {'year': 2023, 'loss_area_ha': 1523.4}
            ],
            'metadata': {
                'source': 'Global Forest Watch',
                'dataset': 'University of Maryland Tree Cover Loss'
            }
        }
        # Temporarily remove charts to test basic export
        # 'charts': [
        #     {
        #         'type': 'line',
        #         'title': 'Forest Loss Trend (2020-2023)',
        #         'x_data': [2020, 2021, 2022, 2023],
        #         'y_data': [1250.5, 1387.2, 1456.8, 1523.4],
        #         'x_label': 'Year',
        #         'y_label': 'Forest Loss (hectares)'
        #     },
        #     {
        #         'type': 'bar',
        #         'title': 'Annual Forest Loss by Year',
        #         'x_data': ['2020', '2021', '2022', '2023'],
        #         'y_data': [1250.5, 1387.2, 1456.8, 1523.4],
        #         'x_label': 'Year',
        #         'y_label': 'Forest Loss (hectares)'
        #     }
        # ]
    }

def example_usage():
    """Example of how to use the report exporter."""
    exporter = ReportExporter()
    sample_data = create_sample_report_data()
    
    try:
        # Export to PDF
        pdf_path = exporter.export_report(sample_data, 'pdf')
        print(f"PDF report exported to: {pdf_path}")
        
        # Export to Word
        word_path = exporter.export_report(sample_data, 'docx')
        print(f"Word report exported to: {word_path}")
        
    except ImportError as e:
        print(f"Import error: {e}")
        print("Install required packages: pip install reportlab python-docx matplotlib")
    except Exception as e:
        print(f"Export error: {e}")

if __name__ == "__main__":
    example_usage()